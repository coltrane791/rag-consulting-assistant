# rag_chat.py
import os, argparse, json, time, re
from datetime import datetime
from dotenv import load_dotenv
import openai

from vector_store import load_index
from retriever import retrieve_with_expansion_and_rerank
from rag_report import build_prompt, ask_llm, _label_for_chunk  # reuse your existing prompt + call

DRAFT_PATH = "draft.md"
SECTIONS = ["Title", "Executive Summary", "Findings", "Evidence", "Recommendations"]

# -------------- Draft helpers --------------
def ensure_draft():
    if not os.path.exists(DRAFT_PATH):
        with open(DRAFT_PATH, "w", encoding="utf-8") as f:
            f.write("# Title\n\n"
                    "# Executive Summary\n\n"
                    "# Findings\n\n"
                    "# Evidence\n\n"
                    "# Recommendations\n")
    return DRAFT_PATH

def read_draft():
    with open(DRAFT_PATH, "r", encoding="utf-8") as f:
        return f.read()

def write_draft(text):
    with open(DRAFT_PATH, "w", encoding="utf-8") as f:
        f.write(text)

def get_section_text(section: str) -> str:
    """Return the text under '# <Section>' until the next '# ' or EOF."""
    text = read_draft()
    pattern = rf"(?ms)^#\s+{re.escape(section)}\s*\n(.*?)(?=^\#\s+|\Z)"
    m = re.search(pattern, text)
    return (m.group(1).strip() if m else "").strip()

def insert_into_section(action: str, section: str, content: str):
    """Append or replace content in a markdown H1 section."""
    text = read_draft()
    if section not in SECTIONS:
        print(f"Unknown section. Choose from: {', '.join(SECTIONS)}")
        return
    pattern = rf"(?ms)^(\#\s+{re.escape(section)}\s*\n)(.*?)(?=^\#\s+|\Z)"
    m = re.search(pattern, text)
    if not m:
        # create section if missing
        if not text.endswith("\n"):
            text += "\n"
        block = f"# {section}\n\n{content.strip()}\n"
        write_draft(text + block)
        return
    start_hdr, body = m.group(1), m.group(2)
    if action == "append":
        new_body = (body.rstrip() + "\n\n" + content.strip() + "\n").lstrip()
    elif action == "replace":
        new_body = ("\n" + content.strip() + "\n")
    else:
        return
    new_text = text[:m.start()] + start_hdr + new_body + text[m.end():]
    write_draft(new_text)

def save_session_line(path, obj):
    with open(path, "a", encoding="utf-8") as f:
        f.write(json.dumps(obj, ensure_ascii=False) + "\n")

# -------------- Chat loop --------------
def main():
    load_dotenv()
    openai.api_key = os.getenv("OPENAI_API_KEY")

    parser = argparse.ArgumentParser(description="Interactive RAG chat over your corpus")
    parser.add_argument("--session", type=str, default=datetime.now().strftime("%Y%m%d-%H%M%S"),
                        help="Session name for saving chat logs (jsonl).")
    parser.add_argument("--top-k", type=int, default=4)
    parser.add_argument("--expand-queries", type=int, default=2)
    parser.add_argument("--per-query-k", type=int, default=6)
    parser.add_argument("--no-llm-rerank", action="store_true")
    parser.add_argument("--rerank-model", type=str, default="gpt-4o-mini")
    parser.add_argument("-g", "--use-general-knowledge", action="store_true")
    parser.add_argument("--cite", action="store_true")
    parser.add_argument("--model", type=str, default="gpt-4")
    parser.add_argument("--temperature", type=float, default=None)
    args = parser.parse_args()

    # Load existing index
    index_dir = "index_store"
    try:
        index, chunks_meta, manifest = load_index(index_dir)
        print(f"Loaded index with {manifest.get('num_vectors')} vectors; model={manifest.get('model_name')}")
    except Exception as e:
        print("No persisted index found. Build it first with rag_report.py.")
        return

    ensure_draft()
    os.makedirs("sessions", exist_ok=True)
    session_path = os.path.join("sessions", f"{args.session}.jsonl")
    print(f"Session log: {session_path}")
    print("Type a question, or commands:")
    print("  /append <Section>      - append last answer to section")
    print("  /replace <Section>     - replace section with last answer")
    print("  /revise <Section>: <instructions>  - rewrite a section with your guidance (and fresh retrieval)")
    print("  /show                  - show current draft")
    print("  /export                - export draft.md to Word (.docx)")
    print("  /quit                  - exit")

    # temperature defaults mirror rag_report.py
    def pick_temp(use_gk: bool, cite: bool, override):
        if override is not None:
            return override
        if cite and use_gk:
            return 0.50
        if cite and not use_gk:
            return 0.35
        if not cite and use_gk:
            return 0.60
        return 0.40

    last_answer = None
    last_chunks = None

    while True:
        try:
            q = input("\n❓> ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\nExiting.")
            break

        if not q:
            continue

        # ---- Commands ----
        if q.lower().startswith("/quit"):
            print("Bye!")
            break

        if q.lower().startswith("/show"):
            print("\n--- draft.md ---\n")
            print(read_draft())
            continue

        if q.lower().startswith("/append "):
            if not last_answer:
                print("No last answer to append. Ask a question first.")
                continue
            section = q.split(" ", 1)[1].strip()
            insert_into_section("append", section, last_answer)
            print(f"Appended to '{section}'.")
            continue

        if q.lower().startswith("/replace "):
            if not last_answer:
                print("No last answer to use. Ask a question first.")
                continue
            section = q.split(" ", 1)[1].strip()
            insert_into_section("replace", section, last_answer)
            print(f"Replaced '{section}'.")
            continue

        if q.lower().startswith("/export"):
            try:
                from docx import Document
                from docx.shared import Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                doc = Document()
                title = doc.add_paragraph("Consulting Analysis Report (Draft)")
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title.runs[0].font.size = Pt(20)
                doc.add_paragraph(datetime.now().strftime("%B %d, %Y")).alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph("")

                for line in read_draft().splitlines():
                    if line.startswith("# "):
                        p = doc.add_paragraph(line[2:])
                        p.runs[0].bold = True
                    else:
                        if line.strip():
                            doc.add_paragraph(line)
                        else:
                            doc.add_paragraph("")
                os.makedirs("output", exist_ok=True)
                out = os.path.join("output", f"draft_export_{int(time.time())}.docx")
                doc.save(out)
                print(f"Exported to {out}")
            except Exception as e:
                print(f"Export failed: {e}")
            continue

        if q.lower().startswith("/note "):
            # /note <Section>: <your text>
            m = re.match(r"(?is)^/note\s+([^:]+)\s*:\s*(.+)$", q)
            if not m:
                print("Usage: /note <Section>: <your text>")
                continue
            section = m.group(1).strip()
            note_text = m.group(2).strip()
            insert_into_section("append", section, note_text)
            print(f"Added your note to '{section}'.")
            # Optionally log it
            log = {
                "ts": datetime.now().isoformat(),
                "command": "note",
                "section": section,
                "note": note_text
            }
            save_session_line(session_path, log)
            continue

        if q.lower().startswith("/revise "):
            # /revise <Section>: <instructions>
            m = re.match(r"(?is)^/revise\s+([^:]+)\s*:\s*(.+)$", q)
            if not m:
                print("Usage: /revise <Section>: <instructions>")
                continue
            section = m.group(1).strip()
            instructions = m.group(2).strip()
            if section not in SECTIONS:
                print(f"Unknown section. Choose from: {', '.join(SECTIONS)}")
                continue
            current_text = get_section_text(section)

            # Build a “revision question” so retrieval can bring in any new evidence
            revision_question = f"Revise the section '{section}' per: {instructions}. Focus on accuracy and clarity."
            candidates = retrieve_with_expansion_and_rerank(
                index,
                chunks_meta,
                revision_question,
                top_k=max(3, min(6, args.top_k + 1)),   # slightly more context for revision
                expand_n=max(1, args.expand_queries),   # ensure at least some expansion
                per_query_k=args.per_query_k,
                use_llm_rerank=(not args.no_llm_rerank),
                rerank_model=args.rerank_model
            )

            # Build a prompt that includes the current section and instructions
            # but still grounds on the retrieved excerpts.
            preface = (
                "You are revising a report section. Use the provided excerpts (and general knowledge if enabled) "
                "to rewrite the section according to the instructions. Keep it well-structured and concise. "
                "Preserve or improve citations if --cite is used in the prompt.\n\n"
                f"Current section name: {section}\n"
                f"Current section text:\n{current_text or '(empty)'}\n\n"
                f"Revision instructions:\n{instructions}\n\n"
            )
            prompt = preface + build_prompt(
                candidates,
                revision_question,
                allow_general_knowledge=args.use_general_knowledge,
                require_citations=args.cite
            )

            temp = pick_temp(args.use_general_knowledge, args.cite, args.temperature)
            print("✏️  Revising section...")
            revised = ask_llm(prompt, model=args.model, temperature=temp)
            insert_into_section("replace", section, revised)
            print(f"Revised '{section}'.")
            # record as last answer so you can /append elsewhere if you like
            last_answer = revised
            last_chunks = candidates

            # Log the revision
            log = {
                "ts": datetime.now().isoformat(),
                "command": "revise",
                "section": section,
                "instructions": instructions,
                "result": revised,
                "chunks": [{"label": _label_for_chunk(c), "source": c["source"], "page": c.get("page"), "chunk_id": c.get("chunk_id")} for c in candidates]
            }
            os.makedirs("sessions", exist_ok=True)
            save_session_line(session_path, log)
            continue

        # ---- Regular question -> retrieve -> answer ----
        candidates = retrieve_with_expansion_and_rerank(
            index,
            chunks_meta,
            q,
            top_k=args.top_k,
            expand_n=args.expand_queries,
            per_query_k=args.per_query_k,
            use_llm_rerank=(not args.no_llm_rerank),
            rerank_model=args.rerank_model
        )

        prompt = build_prompt(
            candidates, q,
            allow_general_knowledge=args.use_general_knowledge,
            require_citations=args.cite
        )
        temp = pick_temp(args.use_general_knowledge, args.cite, args.temperature)

        print("📤 Answering...")
        answer = ask_llm(prompt, model=args.model, temperature=temp)
        print("\n--- Answer ---\n")
        print(answer)

        last_answer = answer
        last_chunks = candidates

        # Log Q&A
        log = {
            "ts": datetime.now().isoformat(),
            "question": q,
            "answer": answer,
            "chunks": [{"label": _label_for_chunk(c), "source": c["source"], "page": c.get("page"), "chunk_id": c.get("chunk_id")} for c in candidates],
            "settings": {
                "top_k": args.top_k,
                "expand_queries": args.expand_queries,
                "per_query_k": args.per_query_k,
                "rerank": not args.no_llm_rerank,
                "model": args.model,
                "general_knowledge": args.use_general_knowledge,
                "cite": args.cite
            }
        }
        save_session_line(session_path, log)

if __name__ == "__main__":
    main()
