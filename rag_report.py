# rag_report.py
import os
import argparse
from datetime import datetime
from dotenv import load_dotenv
import openai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from document_loader import load_and_chunk_folder
from embed_and_search import build_vector_store, search_index, EMBED_MODEL
from vector_store import save_index, load_index, file_sha256
from retriever import retrieve_with_expansion_and_rerank

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# ---------- Helpers ----------
def _label_for_chunk(c):
    # If page is known (PDF), show it; else fall back to chunk id
    if c.get("page"):
        return f"{c['source']} p.{c['page']}"
    return f"{c['source']} #{c['chunk_id']}"

def _labels_for_chunks(chunks):
    return [f"[{_label_for_chunk(c)}]" for c in chunks]

def build_prompt(chunks, question, allow_general_knowledge: bool, require_citations: bool):
    """
    Build the LLM prompt.
    - allow_general_knowledge: Hybrid mode (use excerpts + model knowledge)
    - require_citations: Strict citation mode (inline [Source p.X] after factual claims)
    """
    excerpts_block = "\n\n".join([f"[{_label_for_chunk(c)}] {c['text']}" for c in chunks])
    labels_list = ", ".join(_labels_for_chunks(chunks))

    if allow_general_knowledge:
        instruction = (
            "You are an expert analyst. Using both the provided document excerpts and your own general knowledge, "
            "compare, contrast, and synthesize to answer the user's question. If there is any conflict, prioritize "
            "the document excerpts. If you add anything not in the excerpts, treat it as general background."
        )
    else:
        instruction = (
            "You are an expert analyst. Using only the information in the provided document excerpts, "
            "answer the user's question concisely but with sufficient detail. If information is insufficient, say so."
        )

    if require_citations:
        instruction += (
            "\n\nCITATION RULES:\n"
            f"- After each factual sentence or claim, include an inline citation using one of these labels: {labels_list}.\n"
            "- Place the citation at the end of the sentence in square brackets, e.g., “... statement. [Source p.2]”.\n"
            "- Only use the labels provided above; do not invent new ones. If multiple excerpts support a claim, cite the most direct.\n"
            "- If a claim cannot be supported by the excerpts, state that explicitly (and in Hybrid mode, mark general knowledge as such)."
        )

    return (
        f"{instruction}\n\n"
        f"Document excerpts (each labeled with source):\n{excerpts_block}\n\n"
        f"User question: {question}\n\n"
        "Answer:"
    )

def ask_llm(prompt, *, model="gpt-4", temperature=0.4):
    resp = openai.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a careful, factual consultant."},
            {"role": "user", "content": prompt}
        ],
        temperature=temperature
    )
    return resp.choices[0].message.content.strip()

def _has_any_label(answer_text: str, labels: list[str]) -> bool:
    at_least_one = any(label in answer_text for label in labels)
    return at_least_one

def _append_sources_if_missing(answer_text: str, labels: list[str]) -> str:
    # Safety net: if no inline labels appear, append a compact Sources line.
    if _has_any_label(answer_text, labels):
        return answer_text
    compact = "; ".join(labels)
    return answer_text + f"\n\n(Sources: {compact})"

def write_report_docx(question, answer, chunks, out_path, *, mode_label: str, require_citations: bool):
    doc = Document()

    # Title
    title = doc.add_paragraph("Consulting Analysis Report")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(20)

    date_p = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(11)

    mode_p = doc.add_paragraph(f"Mode: {mode_label}")
    mode_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mode_p.runs[0].font.size = Pt(10)

    if require_citations:
        cite_p = doc.add_paragraph("Citations: Inline labels like [Source p.X] appear after factual sentences.")
        cite_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cite_p.runs[0].font.size = Pt(10)

    doc.add_paragraph("")

    # Question
    h1 = doc.add_paragraph("Question"); h1.runs[0].bold = True
    doc.add_paragraph(question)

    # Answer
    h2 = doc.add_paragraph("Answer"); h2.runs[0].bold = True
    for para in answer.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # Sources appendix
    doc.add_paragraph("")
    h3 = doc.add_paragraph("Sources (Excerpts Used)"); h3.runs[0].bold = True
    for i, ch in enumerate(chunks, start=1):
        p = doc.add_paragraph(f"Excerpt {i} — {_label_for_chunk(ch)}:")
        p.runs[0].bold = True
        doc.add_paragraph(ch["text"])

    doc.save(out_path)

def compute_corpus_sources(folder: str):
    """Return [{path, hash}] for all .docx/.pdf in folder (single-folder mode)."""
    paths = [os.path.join(folder, f) for f in os.listdir(folder)
             if f.lower().endswith(".docx") or f.lower().endswith(".pdf")]
    return [{"path": os.path.abspath(p), "hash": file_sha256(p)} for p in sorted(paths)]

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="RAG report generator")
    parser.add_argument("-g", "--use-general-knowledge", action="store_true",
                        help="Allow the model to use its general knowledge in addition to excerpts (hybrid mode).")
    parser.add_argument("--cite", action="store_true",
                        help="Require inline citations like [Source p.X] after factual claims.")
    parser.add_argument("--top-k", type=int, default=3, help="Number of excerpts to retrieve (default: 3).")
    parser.add_argument("--max-words", type=int, default=300, help="Chunk size in words (default: 300).")
    parser.add_argument("--model", type=str, default="gpt-4", help="OpenAI chat model (default: gpt-4).")
    parser.add_argument("--temperature", type=float, default=None,
                        help="Override temperature. If not set, uses 0.4 (strict), 0.6 (hybrid), 0.35 (strict+cite), 0.5 (hybrid+cite).")
    parser.add_argument("--expand-queries", type=int, default=0,
                    help="Generate N paraphrases of the question for broader recall (default: 0).")
    parser.add_argument("--per-query-k", type=int, default=8,
                    help="Retrieve this many candidates per query/expansion (default: 8).")
    parser.add_argument("--no-llm-rerank", action="store_true",
                    help="Disable LLM reranking; sort by FAISS distance only.")
    parser.add_argument("--rerank-model", type=str, default="gpt-4o-mini",
                    help="Model to use for LLM reranking (default: gpt-4o-mini).")
    args = parser.parse_args()

    corpus_dir = "input"
    index_dir  = "index_store"
    os.makedirs("output", exist_ok=True)

    # Load vs build
    use_existing = input("Use existing index if available? [Y/n]: ").strip().lower()
    load_requested = (use_existing in ("", "y", "yes"))

    index = None
    chunks_meta = None

    if load_requested:
        try:
            index, chunks_meta, manifest = load_index(index_dir)
            # Model + corpus checks
            current_sources = compute_corpus_sources(corpus_dir)
            same_model = (manifest.get("model_name") == EMBED_MODEL)
            same_set = (manifest.get("sources") == current_sources)
            if not same_model:
                print(f"⚠️ Index model {manifest.get('model_name')} != current {EMBED_MODEL}. Rebuilding.")
                index = None
            elif not same_set:
                print("⚠️ Corpus has changed (file added/removed/modified). Rebuilding.")
                index = None
            else:
                print(f"✅ Loaded index ({manifest['num_vectors']} vectors) for {len(current_sources)} doc(s).")
        except Exception as e:
            print(f"ℹ️ Could not load existing index ({e}). Will build a new one.")

    if index is None:
        print("🔄 Loading & chunking all documents in input/ ...")
        chunks_meta = load_and_chunk_folder(corpus_dir, max_words=args.max_words, overlap=60)

        print("🔄 Embedding and indexing...")
        index = build_vector_store(chunks_meta)

        sources = compute_corpus_sources(corpus_dir)
        save_index(
            index=index,
            chunks_meta=chunks_meta,
            index_dir=index_dir,
            model_name=EMBED_MODEL,
            sources=sources,
            max_words=args.max_words
        )
        print(f"✅ Saved index to {index_dir} for {len(sources)} doc(s).")

    # Query
    question = input("❓ What would you like to ask about the corpus? ")

    # Ensure chunks_meta is present if we loaded index
    if chunks_meta is None:
        _, chunks_meta, _ = load_index(index_dir)

    top_chunks = retrieve_with_expansion_and_rerank(
    index,
    chunks_meta,
    question,
    top_k=args.top_k,
    expand_n=args.expand_queries,
    per_query_k=args.per_query_k,
    use_llm_rerank=(not args.no_llm_rerank),
    rerank_model=args.rerank_model
)
    # Prompt + LLM
    prompt = build_prompt(
        top_chunks,
        question,
        allow_general_knowledge=args.use_general_knowledge,
        require_citations=args.cite
    )

    # Temperature defaults tuned for mode
    if args.temperature is not None:
        temp = args.temperature
    else:
        if args.cite and args.use_general_knowledge:
            temp = 0.50
        elif args.cite and not args.use_general_knowledge:
            temp = 0.35
        elif not args.cite and args.use_general_knowledge:
            temp = 0.60
        else:
            temp = 0.40

    print("📤 Querying LLM...")
    answer = ask_llm(prompt, model=args.model, temperature=temp)

    # Safety net: ensure at least one label is present if --cite used
    if args.cite:
        labels = _labels_for_chunks(top_chunks)
        answer = _append_sources_if_missing(answer, labels)

    # Save report
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    out_path = os.path.join("output", f"report_{ts}.docx")
    mode_label = "Hybrid (RAG + General Knowledge)" if args.use_general_knowledge else "Strict RAG (Excerpts Only)"
    write_report_docx(question, answer, top_chunks, out_path, mode_label=mode_label, require_citations=args.cite)
    print(f"✅ Report saved to {out_path}")
